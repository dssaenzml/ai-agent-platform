from typing import Any, Dict

from docx import Document

from ...vector_db.agent_general import kbm
from ..document_gen_tool import DocumentGeneratorTool
from ..tool_wrapper.document_gen_tool_wrapper import DocumentGeneratorToolWrapper


# SoW template function
def consultancy_services_sow_document_template(
    document: Document,
    document_input: Dict[str, Any],
) -> None:
    """
    SoW template function to generate a document.

    :param document: The Document object to be populated.
    :param document_input: A dictionary containing the document content.
    """

    # Function to delete a paragraph
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._element = p._p = None

    # List of paragraph indices to delete
    paragraph_indices_to_delete = [
        32,
        33,
        34,
        35,
        40,
        41,
        42,
        43,
    ]

    # Delete the specified paragraphs
    for index in sorted(paragraph_indices_to_delete, reverse=True):
        delete_paragraph(document.paragraphs[index])

    # 1.0	Preamble
    document.paragraphs[12].text = document_input.get("preamble", "")

    # 2.0	Scope of Services
    ## 2.1	General
    document.paragraphs[15].text = document_input.get("general_sow", "")

    ## 2.2	Description of the Services
    document.paragraphs[17].text = document_input.get(
        "description_of_services",
        "",
    )

    ## 2.3	Codes & Standards
    document.paragraphs[19].text = document_input.get("codes_standards", "")

    ## 2.4	Drawings & Specifications
    document.paragraphs[21].text = document_input.get(
        "drawings_specifications",
        "",
    )

    ## 2.5	Review Periods, Meetings and Reporting Requirements
    document.paragraphs[23].text = document_input.get(
        "review_meetings_reporting",
        "",
    )

    ## 2.6	Training Requirements
    document.paragraphs[25].text = document_input.get(
        "training_requirements",
        "",
    )

    ## 2.7	Interface Requirements
    document.paragraphs[27].text = document_input.get(
        "interface_requirements",
        "",
    )

    ## 2.8	Deliverables
    document.paragraphs[29].text = document_input.get("deliverables", "")

    # 3.0	Exclusions
    document.paragraphs[31].text = document_input.get("exclusions", "")

    # 4.0	Optional Scope
    document.paragraphs[33].text = document_input.get("optional_scope", "")

    # 5.0	Facilities provided by the Employer
    document.paragraphs[35].text = document_input.get(
        "facilities_by_employer",
        "",
    )


consultancy_services_sow_tool_wrapper = DocumentGeneratorToolWrapper(
    document_template=consultancy_services_sow_document_template
)

consultancy_services_sow_tool = DocumentGeneratorTool(
    tool_wrapper=consultancy_services_sow_tool_wrapper,
    kbm=kbm,
)

# from docx import Document

# # Load the document
# doc_path = "app/file_template/procurement-agent/contracts/service_agreement_template.md"
# doc = Document(doc_path)

# document_input = {
#     "preamble": "The project is located in Abu Dhabi and aims to enhance maritime trade. The key parties involved are AD Ports Group and various contractors.",
#     "general": "The general scope of services includes project management, logistics, and coordination with stakeholders.",
#     "description_of_services": "The services to be performed include detailed planning, execution, and monitoring of the project activities.",
#     "codes_standards": "The project will adhere to international maritime codes and standards.",
#     "drawings_specifications": "Detailed drawings and specifications will be provided to ensure compliance with project requirements.",
#     "review_meetings_reporting": "Regular review meetings will be held, and detailed reports will be submitted for approval.",
#     "training_requirements": "Training sessions will be conducted to ensure all personnel are well-versed with the project requirements.",
#     "interface_requirements": "The contractor will manage interfaces with other contractors and ensure smooth execution of services.",
#     "deliverables": "Deliverables include project plans, progress reports, and final project documentation.",
#     "exclusions": "The scope excludes services provided by other contractors and any costs borne by the Employer.",
#     "optional_scope": "Optional scope items include additional training sessions and extended project support.",
#     "facilities_by_employer": "The Employer will provide office space, access to necessary data, and other support services required for the project.",
# }

# # Call the function to find and replace the text
# consultancy_services_sow_document_template(doc, document_input)
# doc.save('dummy_doc_output.docx')
